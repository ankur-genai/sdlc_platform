import os
import sys
import docx
from docx import Document
from docx.shared import Inches, Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.oxml import OxmlElement, parse_xml
from docx.oxml.ns import nsdecls, qn

def set_cell_background(cell, fill_hex):
    """Sets background color of a table cell."""
    tcPr = cell._element.get_or_add_tcPr()
    shd = parse_xml(f'<w:shd {nsdecls("w")} w:fill="{fill_hex}"/>')
    tcPr.append(shd)

def set_cell_margins(cell, top=120, bottom=120, left=150, right=150):
    """Sets internal padding (margins) for a table cell in twips (1 pt = 20 twips)."""
    tcPr = cell._element.get_or_add_tcPr()
    tcMar = OxmlElement('w:tcMar')
    for m, val in [('top', top), ('bottom', bottom), ('left', left), ('right', right)]:
        node = OxmlElement(f'w:{m}')
        node.set(qn('w:w'), str(val))
        node.set(qn('w:type'), 'dxa')
        tcMar.append(node)
    tcPr.append(tcMar)

def add_styled_heading(doc, text, level):
    """Adds a styled heading with custom fonts and colors."""
    h = doc.add_heading(text, level=level)
    h.paragraph_format.space_before = Pt(14)
    h.paragraph_format.space_after = Pt(6)
    h.paragraph_format.keep_with_next = True
    
    run = h.runs[0]
    run.font.name = 'Arial'
    if level == 1:
        run.font.size = Pt(20)
        run.font.bold = True
        run.font.color.rgb = RGBColor(11, 19, 43) # Deep Navy
    elif level == 2:
        run.font.size = Pt(15)
        run.font.bold = True
        run.font.color.rgb = RGBColor(0, 114, 198) # EY Blue
    elif level == 3:
        run.font.size = Pt(12.5)
        run.font.bold = True
        run.font.color.rgb = RGBColor(45, 55, 72) # Slate Dark
    return h

def add_body_paragraph(doc, text, bold_prefix=None, space_after=6):
    """Adds a body paragraph formatted according to EY consulting standards."""
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(0)
    p.paragraph_format.space_after = Pt(space_after)
    p.paragraph_format.line_spacing = 1.15

    if bold_prefix:
        r_pre = p.add_run(bold_prefix)
        r_pre.font.name = 'Calibri'
        r_pre.font.size = Pt(11)
        r_pre.font.bold = True
        r_pre.font.color.rgb = RGBColor(26, 29, 32)

    r_text = p.add_run(text)
    r_text.font.name = 'Calibri'
    r_text.font.size = Pt(11)
    r_text.font.color.rgb = RGBColor(40, 44, 52)
    return p

def add_callout_box(doc, title, text, box_type="NOTE"):
    """Adds a stylish EY callout box for notes, warnings, or key insights."""
    tbl = doc.add_table(rows=1, cols=1)
    tbl.alignment = WD_TABLE_ALIGNMENT.CENTER
    tbl.autofit = False
    
    cell = tbl.cell(0, 0)
    cell.width = Inches(6.5)
    
    if box_type == "NOTE":
        bg_hex = "F0F4F8"
        border_hex = "0072C6"
        icon_str = "ℹ️ EY CONSULTING NOTE"
    elif box_type == "INSIGHT":
        bg_hex = "FFFDF0"
        border_hex = "FFE600"
        icon_str = "💡 ARCHITECTURAL INSIGHT"
    else:
        bg_hex = "F4F6F9"
        border_hex = "1A1D20"
        icon_str = "📌 KEY HIGHLIGHT"

    set_cell_background(cell, bg_hex)
    set_cell_margins(cell, top=140, bottom=140, left=180, right=180)
    
    # Border XML
    tcPr = cell._element.get_or_add_tcPr()
    tcBorders = parse_xml(
        f'<w:tcBorders {nsdecls("w")}>'
        f'<w:left w:val="single" w:sz="24" w:space="0" w:color="{border_hex}"/>'
        f'<w:top w:val="none"/>'
        f'<w:right w:val="none"/>'
        f'<w:bottom w:val="none"/>'
        f'</w:tcBorders>'
    )
    tcPr.append(tcBorders)

    p = cell.paragraphs[0]
    p.paragraph_format.space_after = Pt(4)
    r_title = p.add_run(f"{icon_str}: {title}\n")
    r_title.font.name = 'Arial'
    r_title.font.size = Pt(10.5)
    r_title.font.bold = True
    r_title.font.color.rgb = RGBColor(11, 19, 43)

    r_body = p.add_run(text)
    r_body.font.name = 'Calibri'
    r_body.font.size = Pt(10)
    r_body.font.color.rgb = RGBColor(45, 55, 72)
    
    # Empty line after table for spacing
    p_space = doc.add_paragraph()
    p_space.paragraph_format.space_after = Pt(4)

def add_figure_placeholder(doc, fig_num, title, placeholder_desc, detail_caption):
    """Adds a visual screenshot placeholder with figure caption and description."""
    p_fig = doc.add_paragraph()
    p_fig.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p_fig.paragraph_format.space_before = Pt(10)
    p_fig.paragraph_format.space_after = Pt(4)

    # Box for screenshot
    tbl = doc.add_table(rows=1, cols=1)
    tbl.alignment = WD_TABLE_ALIGNMENT.CENTER
    cell = tbl.cell(0, 0)
    cell.width = Inches(6.5)
    set_cell_background(cell, "F8F9FA")
    set_cell_margins(cell, top=200, bottom=200, left=200, right=200)

    tcPr = cell._element.get_or_add_tcPr()
    tcBorders = parse_xml(
        f'<w:tcBorders {nsdecls("w")}>'
        f'<w:left w:val="single" w:sz="12" w:space="0" w:color="CBD5E1"/>'
        f'<w:top w:val="single" w:sz="12" w:space="0" w:color="CBD5E1"/>'
        f'<w:right w:val="single" w:sz="12" w:space="0" w:color="CBD5E1"/>'
        f'<w:bottom w:val="single" w:sz="12" w:space="0" w:color="CBD5E1"/>'
        f'</w:tcBorders>'
    )
    tcPr.append(tcBorders)

    p_inner = cell.paragraphs[0]
    p_inner.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r_holder = p_inner.add_run(f"📷 [ SCREENSHOT PLACEHOLDER ]\n{placeholder_desc}")
    r_holder.font.name = 'Arial'
    r_holder.font.size = Pt(10)
    r_holder.font.bold = True
    r_holder.font.color.rgb = RGBColor(100, 116, 139)

    # Caption
    p_cap = doc.add_paragraph()
    p_cap.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p_cap.paragraph_format.space_before = Pt(4)
    p_cap.paragraph_format.space_after = Pt(10)
    
    r_num = p_cap.add_run(f"Figure {fig_num} — {title}\n")
    r_num.font.name = 'Arial'
    r_num.font.size = Pt(9.5)
    r_num.font.bold = True
    r_num.font.color.rgb = RGBColor(11, 19, 43)

    r_desc = p_cap.add_run(detail_caption)
    r_desc.font.name = 'Calibri'
    r_desc.font.size = Pt(9.5)
    r_desc.font.italic = True
    r_desc.font.color.rgb = RGBColor(100, 116, 139)

def format_styled_table(table, col_widths, headers, data):
    """Formats a Word table with EY headers, alternating row colors, and proper borders."""
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    table.autofit = False

    # Header Row
    hdr_cells = table.rows[0].cells
    for i, header_text in enumerate(headers):
        hdr_cells[i].text = header_text
        hdr_cells[i].width = col_widths[i]
        set_cell_background(hdr_cells[i], "0B132B") # EY Dark Navy
        set_cell_margins(hdr_cells[i], top=140, bottom=140, left=140, right=140)
        p = hdr_cells[i].paragraphs[0]
        p.alignment = WD_ALIGN_PARAGRAPH.LEFT
        for r in p.runs:
            r.font.name = 'Arial'
            r.font.size = Pt(10)
            r.font.bold = True
            r.font.color.rgb = RGBColor(255, 230, 0) # EY Gold

    # Data Rows
    for r_idx, row_data in enumerate(data):
        row_cells = table.add_row().cells
        bg_color = "F4F6F9" if r_idx % 2 == 1 else "FFFFFF"
        for c_idx, cell_value in enumerate(row_data):
            row_cells[c_idx].text = str(cell_value)
            row_cells[c_idx].width = col_widths[c_idx]
            set_cell_background(row_cells[c_idx], bg_color)
            set_cell_margins(row_cells[c_idx], top=100, bottom=100, left=120, right=120)
            p = row_cells[c_idx].paragraphs[0]
            p.alignment = WD_ALIGN_PARAGRAPH.LEFT
            for r in p.runs:
                r.font.name = 'Calibri'
                r.font.size = Pt(9.5)
                r.font.color.rgb = RGBColor(40, 44, 52)
