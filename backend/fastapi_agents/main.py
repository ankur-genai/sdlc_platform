"""
main.py — FastAPI app: CORS for the React frontend, cookie-based JWT auth,
and every endpoint needed for the Friday demo.

Run:
    uvicorn main:app --reload --port 8000

Endpoints
---------
Auth:
    POST /auth/register
    POST /auth/login        (sets access_token + refresh_token HttpOnly cookies)
    GET  /auth/me
    POST /auth/refresh       (bonus — mints a new access token from the refresh cookie)
    POST /auth/logout        (bonus — clears both cookies)

Project wizard & ingestion:
    POST /projects
    POST /ingestion/upload
    DELETE /projects/{project_id}

Live dashboard:
    GET  /dashboard/timeline

Workspace delivery:
    GET  /projects/{project_id}/deliverables
    GET  /generated_artifacts
"""
import os
import secrets
import uuid
import logging
import shutil
from datetime import datetime, timedelta, timezone
from pathlib import Path

from dotenv import load_dotenv

# Load backend/.env explicitly (by path, not cwd) so GROQ_API_KEY/
# DEFAULT_AZURE_OPENAI_*/etc. are always picked up regardless of the
# working directory the server was launched from (e.g. `uvicorn
# fastapi_agents.main:app --app-dir backend` runs with cwd at the repo
# root, not backend/) — this is what makes ".env-only" provider swaps
# (Groq <-> Azure) actually take effect without any code change.
load_dotenv(Path(__file__).resolve().parent.parent / ".env")

# Uvicorn only configures its own "uvicorn"/"uvicorn.access"/"uvicorn.error"
# loggers — the root logger otherwise has no handler, so every plain
# `logging.getLogger(...).info(...)` call in this codebase (e.g.
# llm_service.py's "request served by <provider>" line, added specifically
# so provider routing can be verified via the server log) is silently
# dropped. configure_logging() attaches both a console handler and a
# RotatingFileHandler (logs/sdlc.log) to the root logger, so every module's
# logger — however it was obtained — ends up going to both places.
from .logging_config import configure_logging  # noqa: E402

configure_logging()

import jwt
from cryptography.fernet import Fernet
from fastapi import Depends, FastAPI, File, Form, HTTPException, Request, Response, UploadFile, status
from fastapi.middleware.cors import CORSMiddleware
from passlib.context import CryptContext
from pydantic import BaseModel
from sqlalchemy.orm import Session

from .models import (
    AgentRun,
    Base,
    DashboardTimelineResponse,
    DEMO_MODE,
    Document,
    DocumentUploadResponse,
    GeneratedArtifact,
    GeneratedArtifactOut,
    LoginRequest,
    PROVIDER_KEY_ENCRYPTION_KEY,
    Project,
    ProjectCreate,
    ProjectDeliverable,
    ProjectDeliverablesResponse,
    ProjectOut,
    ProjectStatus,
    ProviderConfiguration,
    RunStatus,
    _ensure_provider_configuration_columns,
    _ensure_pipeline_columns,
    TimelineEvent,
    User,
    UserCreate,
    UserOut,
    engine,
    get_db,
)

# ===========================================================================
# Configuration
# ===========================================================================
# No fixed literal fallback here on purpose — a hardcoded default secret
# baked into source is readable by anyone with the code. If these aren't set
# in .env, generate a random ephemeral one (logged loudly below) so the
# server still starts, but every restart invalidates existing sessions/
# encrypted BYOK rows rather than silently using one fixed, guessable value.
JWT_SECRET_KEY = os.getenv("JWT_SECRET_KEY") or secrets.token_urlsafe(48)
JWT_REFRESH_SECRET_KEY = os.getenv("JWT_REFRESH_SECRET_KEY") or secrets.token_urlsafe(48)
JWT_ALGORITHM = "HS256"
ACCESS_TOKEN_EXPIRE_MINUTES = int(os.getenv("ACCESS_TOKEN_EXPIRE_MINUTES", "30"))
REFRESH_TOKEN_EXPIRE_DAYS = int(os.getenv("REFRESH_TOKEN_EXPIRE_DAYS", "7"))

# PROVIDER_KEY_ENCRYPTION_KEY and DEMO_MODE are centralized in models.py
# (imported above) — every module reads the exact same value the exact same
# way, rather than each independently re-parsing the env var (which is how a
# stale/missing value could previously make one module disagree with another).
if not os.getenv("PROVIDER_KEY_ENCRYPTION_KEY") or not os.getenv("JWT_SECRET_KEY"):
    logging.warning(
        "JWT_SECRET_KEY / PROVIDER_KEY_ENCRYPTION_KEY not set in .env — using a random "
        "ephemeral value for this process. Sessions won't survive a restart and any "
        "already-encrypted BYOK provider keys won't decrypt. Set these in .env for production."
    )
_fernet = Fernet(PROVIDER_KEY_ENCRYPTION_KEY.encode())

COOKIE_SECURE = os.getenv("COOKIE_SECURE", "false").lower() == "true"
STORAGE_BASE_PATH = Path(os.getenv("STORAGE_BASE_PATH", "./storage"))
DEMO_EMAIL = "ishratbhullar@gmail.com"
logger = logging.getLogger("sdlc.demo")

CORS_ALLOWED_ORIGINS = [
    origin.strip()
    for origin in os.getenv(
        "FRONTEND_URLS",
        "http://localhost:5173,http://127.0.0.1:5173,http://localhost:3000,http://127.0.0.1:3000",
    ).split(",")
    if origin.strip()
]

ACCESS_COOKIE_NAME = "access_token"
REFRESH_COOKIE_NAME = "refresh_token"

def _extract_document_text(path) -> str:
    """Extract plain text from an uploaded BRD/RFP document (PDF or DOCX) so
    ingestion actually grounds the pipeline in the real document content.
    Falls back gracefully (empty string) for unsupported/unreadable files —
    upload still succeeds, it just won't enrich the project description."""
    from pathlib import Path as _Path
    p = _Path(path)
    suffix = p.suffix.lower()
    try:
        if suffix == ".pdf":
            try:
                import pdfplumber
                with pdfplumber.open(str(p)) as pdf:
                    return "\n\n".join((page.extract_text() or "") for page in pdf.pages)
            except Exception:
                pass
            try:
                import pypdf
                reader = pypdf.PdfReader(str(p))
                return "\n\n".join((page.extract_text() or "") for page in reader.pages)
            except Exception:
                pass
        if suffix in (".docx", ".doc"):
            import docx  # python-docx
            d = docx.Document(str(p))
            return "\n".join(para.text for para in d.paragraphs)
        if suffix in (".txt", ".md"):
            return p.read_text(encoding="utf-8", errors="ignore")
    except Exception as exc:
        logger.warning("[Ingestion] _extract_document_text failed for %s: %s", p.name, exc)
    return ""


pwd_context = CryptContext(schemes=["pbkdf2_sha256"], deprecated="auto")
# In DEMO_MODE, all paths bypass auth — the sentinel value "*" is checked below.
DEMO_AUTH_BYPASS_ALL = True  # set to False to require login even in demo

# ===========================================================================
# App + CORS
# ===========================================================================
app = FastAPI(title="EY Autonomous SDLC Studio")

app.add_middleware(
    CORSMiddleware,
    allow_origins=CORS_ALLOWED_ORIGINS,
    allow_credentials=True,
    allow_methods=["GET", "POST", "PUT", "PATCH", "DELETE", "OPTIONS"],
    allow_headers=["*"],
)

@app.middleware("http")
async def strip_api_prefix(request: Request, call_next):
    if request.url.path.startswith("/api/"):
        request.scope["path"] = request.url.path[4:]
    return await call_next(request)


@app.on_event("startup")
def on_startup() -> None:
    Base.metadata.create_all(bind=engine)
    _ensure_provider_configuration_columns(engine)
    _ensure_pipeline_columns(engine)
    STORAGE_BASE_PATH.mkdir(parents=True, exist_ok=True)


@app.get("/health", tags=["meta"])
def health() -> dict:
    return {"status": "ok"}


# ===========================================================================
# Auth helpers
# ===========================================================================
def hash_password(plain_password: str) -> str:
    return pwd_context.hash(plain_password)


def verify_password(plain_password: str, hashed_password: str) -> bool:
    return pwd_context.verify(plain_password, hashed_password)


def _create_token(user_id: int, token_type: str, secret: str, expires_delta: timedelta, remember: bool = True) -> str:
    now = datetime.now(timezone.utc)
    payload = {"sub": str(user_id), "type": token_type, "iat": now, "exp": now + expires_delta, "remember": remember}
    return jwt.encode(payload, secret, algorithm=JWT_ALGORITHM)


def create_access_token(user_id: int, remember: bool = True) -> str:
    return _create_token(user_id, "access", JWT_SECRET_KEY, timedelta(minutes=ACCESS_TOKEN_EXPIRE_MINUTES), remember)


def create_refresh_token(user_id: int, remember: bool = True) -> str:
    return _create_token(user_id, "refresh", JWT_REFRESH_SECRET_KEY, timedelta(days=REFRESH_TOKEN_EXPIRE_DAYS), remember)


def _decode_token(token: str, secret: str, expected_type: str) -> dict:
    try:
        payload = jwt.decode(token, secret, algorithms=[JWT_ALGORITHM])
    except jwt.ExpiredSignatureError:
        raise HTTPException(status.HTTP_401_UNAUTHORIZED, "Token has expired")
    except jwt.InvalidTokenError:
        raise HTTPException(status.HTTP_401_UNAUTHORIZED, "Invalid token")
    if payload.get("type") != expected_type:
        raise HTTPException(status.HTTP_401_UNAUTHORIZED, f"Expected a {expected_type} token")
    return payload


def _set_auth_cookies(response: Response, user_id: int, remember: bool = True) -> None:
    # "Remember me" controls cookie persistence: when remember is True the
    # cookies carry a max_age so they survive browser restarts; when False they
    # are session cookies (max_age omitted) that the browser drops on close, so
    # the user must sign in again next session. The JWT lifetimes are unchanged.
    access_max_age = ACCESS_TOKEN_EXPIRE_MINUTES * 60 if remember else None
    refresh_max_age = REFRESH_TOKEN_EXPIRE_DAYS * 24 * 60 * 60 if remember else None
    response.set_cookie(
        key=ACCESS_COOKIE_NAME,
        value=create_access_token(user_id, remember),
        httponly=True,
        secure=COOKIE_SECURE,
        samesite="lax",
        max_age=access_max_age,
        path="/",
    )
    response.set_cookie(
        key=REFRESH_COOKIE_NAME,
        value=create_refresh_token(user_id, remember),
        httponly=True,
        secure=COOKIE_SECURE,
        samesite="lax",
        max_age=refresh_max_age,
        path="/",
    )


def _get_or_create_demo_user(db: Session) -> User:
    user = db.query(User).filter(User.email == DEMO_EMAIL).first()
    if user is None:
        user = User(
            email=DEMO_EMAIL,
            full_name="Demo User",
            role="developer",
            hashed_password=hash_password("DemoPassword123!"),
        )
        db.add(user)
        db.commit()
        db.refresh(user)
    return user


def get_current_user(request: Request, db: Session = Depends(get_db)) -> User:
    # In DEMO_MODE, bypass auth entirely for all paths
    if DEMO_MODE and DEMO_AUTH_BYPASS_ALL:
        token = request.cookies.get(ACCESS_COOKIE_NAME)
        if token:
            try:
                payload = _decode_token(token, JWT_SECRET_KEY, expected_type="access")
                user = db.get(User, int(payload["sub"]))
                if user:
                    return user
            except Exception:
                pass
        logger.info("DEMO_AUTH_BYPASS path=%s reason=demo_mode email=%s", request.url.path, DEMO_EMAIL)
        return _get_or_create_demo_user(db)

    token = request.cookies.get(ACCESS_COOKIE_NAME)
    if not token:
        logger.info("AUTH_REQUIRED path=%s demo_mode=%s reason=missing_cookie", request.url.path, DEMO_MODE)
        raise HTTPException(status.HTTP_401_UNAUTHORIZED, "Not authenticated")
    try:
        payload = _decode_token(token, JWT_SECRET_KEY, expected_type="access")
    except HTTPException:
        raise
    user = db.get(User, int(payload["sub"]))
    if user is None:
        raise HTTPException(status.HTTP_401_UNAUTHORIZED, "User no longer exists")
    return user


def encrypt_secret(plaintext: str) -> str:
    return _fernet.encrypt(plaintext.encode()).decode()


def _get_project_or_404(db: Session, project_id: int) -> Project:
    project = db.get(Project, project_id)
    if project is None:
        raise HTTPException(status.HTTP_404_NOT_FOUND, f"Project {project_id} not found")
    return project


# ===========================================================================
# Auth endpoints
# ===========================================================================
@app.post("/auth/register", response_model=UserOut, status_code=status.HTTP_201_CREATED, tags=["auth"])
def register(payload: UserCreate, db: Session = Depends(get_db)) -> User:
    if db.query(User).filter(User.email == payload.email).first():
        raise HTTPException(status.HTTP_400_BAD_REQUEST, "A user with that email already exists")
    user = User(
        email=payload.email,
        full_name=payload.full_name,
        role=payload.role.value,
        hashed_password=hash_password(payload.password),
    )
    # Store selected agents for this user (if provided) as a generated project to tie into existing schema.
    # For now we log this association and leave persistence to the project/agent pipeline.
    if payload.selected_agents:
        logging.info("USER_REGISTER_SELECTED_AGENTS email=%s agents=%s", payload.email, payload.selected_agents)
    db.add(user)
    db.commit()
    db.refresh(user)
    return user


@app.post("/auth/login", response_model=UserOut, tags=["auth"])
def login(payload: LoginRequest, response: Response, db: Session = Depends(get_db)) -> User:
    if DEMO_MODE:
        email = payload.email.strip().lower()
        user = db.query(User).filter(User.email == email).first()
        if user is None:
            user = User(
                email=email,
                full_name=email.split("@")[0].capitalize() or "Demo User",
                role="developer",
                hashed_password=hash_password(payload.password or "DemoPassword123!"),
            )
            db.add(user)
            db.commit()
            db.refresh(user)
        _set_auth_cookies(response, user.id, remember=payload.remember_me)
        return user

    user = db.query(User).filter(User.email == payload.email).first()
    if user is None or not verify_password(payload.password, user.hashed_password):
        raise HTTPException(status.HTTP_401_UNAUTHORIZED, "Incorrect email or password")
    _set_auth_cookies(response, user.id, remember=payload.remember_me)
    return user


@app.get("/auth/me", response_model=UserOut, tags=["auth"])
def me(current_user: User = Depends(get_current_user)) -> User:
    return current_user


@app.post("/auth/refresh", tags=["auth"])
def refresh(request: Request, response: Response, db: Session = Depends(get_db)) -> dict:
    token = request.cookies.get(REFRESH_COOKIE_NAME)
    if not token:
        raise HTTPException(status.HTTP_401_UNAUTHORIZED, "No refresh token provided")
    payload = _decode_token(token, JWT_REFRESH_SECRET_KEY, expected_type="refresh")
    user_id = int(payload["sub"])
    if db.get(User, user_id) is None:
        raise HTTPException(status.HTTP_401_UNAUTHORIZED, "User no longer exists")
    remember = bool(payload.get("remember", True))
    response.set_cookie(
        key=ACCESS_COOKIE_NAME,
        value=create_access_token(user_id, remember),
        httponly=True,
        secure=COOKIE_SECURE,
        samesite="lax",
        max_age=ACCESS_TOKEN_EXPIRE_MINUTES * 60 if remember else None,
        path="/",
    )
    return {"detail": "Access token refreshed"}


@app.post("/auth/logout", tags=["auth"])
def logout(response: Response) -> dict:
    response.delete_cookie(ACCESS_COOKIE_NAME, path="/")
    response.delete_cookie(REFRESH_COOKIE_NAME, path="/")
    return {"detail": "Logged out"}


# ===========================================================================
# Project wizard & document ingestion
# ===========================================================================
@app.post("/projects", response_model=ProjectOut, status_code=status.HTTP_201_CREATED, tags=["projects"])
def create_project(
    payload: ProjectCreate,
    db: Session = Depends(get_db),
    current_user: User = Depends(get_current_user),
) -> Project:
    logger.info(
        "PROJECT_CREATE_REQUEST user=%s demo_mode=%s name=%s execution_mode=%s deliverables=%s providers=%s",
        current_user.email,
        DEMO_MODE,
        payload.project_name,
        payload.execution_mode.value,
        [deliverable.value for deliverable in payload.deliverables],
        [provider.value for provider in payload.providers.keys()],
    )
    project = Project(
        name=payload.project_name,
        description=payload.description,
        project_type=payload.project_type.value,
        execution_mode=payload.execution_mode.value,
        build_type=payload.build_type.value,
        status="in_progress",
        project_type_key=payload.project_type_key,
        launch_mode=payload.launch_mode or payload.execution_mode.value,
        selected_agents=payload.selected_agents or None,
        frontend_framework=payload.frontend_framework,
    )
    db.add(project)
    db.flush()

    for deliverable_type in payload.deliverables:
        db.add(ProjectDeliverable(
            project_id=project.id,
            deliverable_type=deliverable_type.value,
            selected=True,
            status=RunStatus.PENDING.value,
        ))

    for provider_name, raw_key in payload.providers.items():
        db.add(ProviderConfiguration(
            project_id=project.id,
            provider_name=provider_name.value,
            enabled=True,
            encrypted_key=encrypt_secret(raw_key),
        ))

    db.add(TimelineEvent(project_id=project.id, stage="Project Created", status=RunStatus.COMPLETED.value))
    db.add(TimelineEvent(
        project_id=project.id,
        stage=f"Launch Mode: {payload.launch_mode or payload.execution_mode.value}",
        status=RunStatus.COMPLETED.value,
    ))
    if payload.manual_stages:
        db.add(TimelineEvent(
            project_id=project.id,
            stage=f"Manual Stages: {', '.join(payload.manual_stages)}",
            status=RunStatus.COMPLETED.value,
        ))
    if payload.build_profile:
        db.add(TimelineEvent(
            project_id=project.id,
            stage=f"Build Profile: {payload.build_profile}",
            status=RunStatus.COMPLETED.value,
        ))
    if payload.providers:
        db.add(TimelineEvent(
            project_id=project.id,
            stage=f"BYOK Providers: {', '.join(provider.value for provider in payload.providers.keys())}",
            status=RunStatus.COMPLETED.value,
        ))

    db.commit()
    db.refresh(project)
    logger.info("PROJECT_CREATE_SUCCESS project_id=%s user=%s status=201", project.id, current_user.email)
    return project


def _generate_project_goal(document_text: str) -> str:
    """Turn raw requirements-document text into a concise 'Project Goal'
    formatted description via the LLM, with a plain-text fallback if the
    model call fails."""
    system = (
        "You are a senior business analyst. Read the provided requirements "
        "document and write a clear, well-structured project description in "
        "'Project Goal' format. Use these sections with markdown headings:\n"
        "**Project Goal** - one or two sentences describing the primary "
        "objective.\n"
        "**Scope** - 3-6 concise bullet points of what will be built.\n"
        "**Key Features** - 3-6 bullet points of the main capabilities.\n"
        "**Target Users** - who will use it.\n"
        "Keep it under 250 words. Do not invent unrelated features; base "
        "everything strictly on the document."
    )
    prompt = f"Requirements document:\n\n{document_text[:12000]}"
    try:
        from .agents.llm_service import LLMService

        result = LLMService().generate_text(system, prompt, temperature=0.3).strip()
        if result:
            return result
    except Exception as exc:
        logger.warning("[Ingestion] Project goal generation failed: %s", exc)
    # Fallback: a trimmed excerpt so the user still gets something usable.
    excerpt = " ".join(document_text.split())[:800]
    return f"**Project Goal**\n\n{excerpt}"


@app.post("/ingestion/analyze-document", tags=["ingestion"])
def analyze_document(
    file: UploadFile = File(...),
    current_user: User = Depends(get_current_user),
) -> dict:
    """Extract text from an uploaded requirements document (PDF/DOCX/TXT/MD)
    and generate a concise, structured 'Project Goal' description from it.

    Used by the New Project wizard's first step, BEFORE a project row exists,
    so this endpoint deliberately does NOT require or create a project. The
    generated text is returned to the client to pre-fill the description
    field; nothing is persisted here."""
    tmp_dir = STORAGE_BASE_PATH / "_uploads_scratch"
    tmp_dir.mkdir(parents=True, exist_ok=True)
    safe_name = f"{uuid.uuid4().hex[:8]}_{file.filename}"
    destination = tmp_dir / safe_name
    try:
        with destination.open("wb") as out_file:
            out_file.write(file.file.read())

        extracted = _extract_document_text(destination).strip()
        if not extracted:
            raise HTTPException(
                status.HTTP_422_UNPROCESSABLE_ENTITY,
                "Could not extract any text from the uploaded document.",
            )

        generated = _generate_project_goal(extracted)
        return {
            "description": generated,
            "extracted_chars": len(extracted),
            "file_name": file.filename,
        }
    finally:
        try:
            destination.unlink(missing_ok=True)
        except Exception:
            pass


@app.post("/ingestion/upload", response_model=DocumentUploadResponse, tags=["ingestion"])
def upload_document(
    project_id: int = Form(..., description="Required so the document can be attached to a project"),
    document_type: str = Form(...),
    file: UploadFile = File(...),
    db: Session = Depends(get_db),
    current_user: User = Depends(get_current_user),
) -> dict:
    project = db.get(Project, project_id)
    if project is None:
        raise HTTPException(status.HTTP_404_NOT_FOUND, "Project not found")

    project_dir = STORAGE_BASE_PATH / f"project_{project_id}"
    project_dir.mkdir(parents=True, exist_ok=True)
    safe_name = f"{uuid.uuid4().hex[:8]}_{file.filename}"
    destination = project_dir / safe_name
    with destination.open("wb") as out_file:
        out_file.write(file.file.read())

    document = Document(
        project_id=project_id,
        file_name=file.filename,
        document_type=document_type,
        storage_path=str(destination),
    )
    db.add(document)
    db.add(TimelineEvent(
        project_id=project_id,
        stage=f"Document Uploaded: {file.filename}",
        status=RunStatus.COMPLETED.value,
    ))

    # Extract real document text (PDF/DOCX) and fold it into the project
    # description so the agent pipeline is actually grounded in the uploaded
    # BRD/RFP content, not just whatever free-text the user typed at creation.
    extracted_chars = 0
    try:
        extracted = _extract_document_text(destination)
        if extracted.strip():
            extracted_chars = len(extracted)
            existing = (project.description or "").strip()
            marker = f"\n\n--- Uploaded document: {file.filename} ---\n"
            project.description = (existing + marker + extracted[:12000]).strip()
            db.add(TimelineEvent(
                project_id=project_id,
                stage=f"Document Ingested: {file.filename} ({extracted_chars} chars extracted)",
                status=RunStatus.COMPLETED.value,
            ))
    except Exception as exc:
        logger.warning("[Ingestion] Text extraction failed for %s: %s", file.filename, exc)

    db.commit()
    db.refresh(document)

    return {"document_id": document.id, "upload_status": "success", "project_reference": project_id,
            "extracted_chars": extracted_chars}


# ===========================================================================
# Live dashboard timeline
# ===========================================================================
@app.get("/dashboard/timeline", response_model=DashboardTimelineResponse, tags=["dashboard"])
def get_dashboard_timeline(
    project_id: int | None = None,
    db: Session = Depends(get_db),
    current_user: User = Depends(get_current_user),
) -> dict:
    project = (
        db.get(Project, project_id)
        if project_id is not None
        else db.query(Project).order_by(Project.created_at.desc()).first()
    )
    if project is None:
        raise HTTPException(status.HTTP_404_NOT_FOUND, "No project found")

    timeline_events = (
        db.query(TimelineEvent)
        .filter(TimelineEvent.project_id == project.id)
        .order_by(TimelineEvent.timestamp.asc())
        .all()
    )
    agent_runs = (
        db.query(AgentRun)
        .filter(AgentRun.project_id == project.id)
        .order_by(AgentRun.id.asc())
        .all()
    )

    return {
        "project_id": project.id,
        "project_status": project.status,
        "timeline_events": timeline_events,
        "agent_runs": agent_runs,
    }


# ===========================================================================
# Workspace delivery
# ===========================================================================
@app.get("/projects/{project_id}/deliverables", response_model=ProjectDeliverablesResponse, tags=["projects"])
def get_project_deliverables(
    project_id: int,
    db: Session = Depends(get_db),
    current_user: User = Depends(get_current_user),
) -> dict:
    project = db.get(Project, project_id)
    if project is None:
        raise HTTPException(status.HTTP_404_NOT_FOUND, "Project not found")

    deliverables = db.query(ProjectDeliverable).filter(ProjectDeliverable.project_id == project_id).all()
    artifacts = (
        db.query(GeneratedArtifact)
        .filter(GeneratedArtifact.project_id == project_id)
        .order_by(GeneratedArtifact.created_at.asc())
        .all()
    )
    return {"project_id": project_id, "deliverables": deliverables, "artifacts": artifacts}


@app.get("/generated_artifacts", response_model=list[GeneratedArtifactOut], tags=["projects"])
def list_generated_artifacts(
    project_id: int | None = None,
    artifact_type: str | None = None,
    db: Session = Depends(get_db),
    current_user: User = Depends(get_current_user),
) -> list[GeneratedArtifact]:
    query = db.query(GeneratedArtifact)
    if project_id is not None:
        query = query.filter(GeneratedArtifact.project_id == project_id)
    if artifact_type is not None:
        query = query.filter(GeneratedArtifact.artifact_type == artifact_type)
    return query.order_by(GeneratedArtifact.created_at.asc()).all()


class BuildStartRequest(BaseModel):
    project_id: int


@app.post("/build/start", tags=["agents"])
async def start_build(
    payload: BuildStartRequest,
    db: Session = Depends(get_db),
    current_user: User = Depends(get_current_user),
) -> dict:
    project = db.get(Project, payload.project_id)
    if project is None:
        raise HTTPException(status.HTTP_404_NOT_FOUND, "Project not found")

    logger.info(
        "BUILD_START_REQUEST project_id=%s user=%s demo_mode=%s",
        payload.project_id,
        current_user.email,
        DEMO_MODE,
    )

    from . import agent_runner

    runs = agent_runner.ensure_agent_runs_exist(db, payload.project_id)

    project.status = ProjectStatus.IN_PROGRESS.value
    db.add(TimelineEvent(project_id=payload.project_id, stage="Autonomous Build Started", status=RunStatus.RUNNING.value))
    if DEMO_MODE:
        db.add(TimelineEvent(
            project_id=payload.project_id,
            stage="Demo Build Queued",
            status=RunStatus.COMPLETED.value,
        ))
        db.commit()
        import asyncio
        asyncio.create_task(agent_runner.run_pipeline(payload.project_id))
        logger.info(
            "BUILD_START_SUCCESS project_id=%s agents_queued=%s demo_mode=%s status=200",
            payload.project_id,
            len(runs),
            DEMO_MODE,
        )
        return {
            "project_id": payload.project_id,
            "agents_queued": len(runs),
            "status": "started",
            "message": f"Demo autonomous build started for project {payload.project_id}",
        }

    db.commit()

    import asyncio
    asyncio.create_task(agent_runner.run_pipeline(payload.project_id))
    logger.info(
        "BUILD_START_SUCCESS project_id=%s agents_queued=%s demo_mode=%s status=200",
        payload.project_id,
        len(runs),
        DEMO_MODE,
    )

    return {
        "project_id": payload.project_id,
        "agents_queued": len(runs),
        "status": "started",
        "message": f"Autonomous build started for project {payload.project_id}",
    }


@app.delete("/projects/{project_id}", status_code=status.HTTP_200_OK, tags=["projects"])
def delete_project(
    project_id: int,
    db: Session = Depends(get_db),
    current_user: User = Depends(get_current_user),
) -> dict:
    """
    Delete a project and all related records.
    Cascading deletes handle agent_runs, approvals, artifacts, timeline_events,
    documents, deliverables, provider_configurations, and review_results.
    The project's uploaded files on disk (STORAGE_BASE_PATH/project_<id>/) are
    removed as well so no orphaned files are left behind.
    """
    project = _get_project_or_404(db, project_id)

    logger.info(
        "PROJECT_DELETE_REQUEST project_id=%s name=%s user=%s",
        project_id, project.name, current_user.email,
    )

    # Collect any explicit document storage paths before the DB rows are gone,
    # then remove the whole per-project storage directory below.
    document_paths = [doc.storage_path for doc in project.documents if doc.storage_path]

    db.delete(project)
    db.commit()

    # Best-effort on-disk cleanup — never let a filesystem hiccup fail the
    # delete after the DB rows are already committed.
    project_dir = STORAGE_BASE_PATH / f"project_{project_id}"
    try:
        if project_dir.exists():
            shutil.rmtree(project_dir, ignore_errors=True)
        # Remove any stray files stored outside the per-project directory.
        for raw_path in document_paths:
            try:
                fp = Path(raw_path)
                if fp.is_file() and project_dir not in fp.parents:
                    fp.unlink(missing_ok=True)
            except Exception as file_exc:  # pragma: no cover - defensive
                logger.warning("PROJECT_DELETE_FILE_SKIP project_id=%s path=%s err=%s", project_id, raw_path, file_exc)
    except Exception as fs_exc:  # pragma: no cover - defensive
        logger.warning("PROJECT_DELETE_FS_CLEANUP_FAILED project_id=%s err=%s", project_id, fs_exc)

    logger.info("PROJECT_DELETE_SUCCESS project_id=%s", project_id)

    return {
        "project_id": project_id,
        "status": "deleted",
        "message": f"Project '{project.name}' and all related records and files deleted successfully",
    }


from .main_extension import router as ext_router
app.include_router(ext_router)
from . import main_extension



if __name__ == "__main__":
    import uvicorn
    uvicorn.run(app, host="0.0.0.0", port=8008)